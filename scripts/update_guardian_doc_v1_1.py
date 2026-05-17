"""
CepSandik Guardian Secim Sureci dokumanini v1.0 -> v1.1'e yukseltir.

Yapilan degisiklikler:
1. Surum satiri "Surum 1.0" -> "Surum 1.1" + revizyon notu
2. Tablo 1 (N degerleri): N=3 satiri silinir
3. Tablo 7 (Sistem Teknik Filtreleri): Filtre #8 yeni topluluk kurali ile guncellenir
4. Bolum 5.3 "Toplam ceremony suresi asla asilmaz" cumlesi guncellenir
5. Dokuman sonuna "11. v1.1 Revizyonlari" bolumu eklenir
"""
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

DOC_PATH = 'CepSandik_Guardian_Secim_Sureci.docx'

doc = Document(DOC_PATH)


def replace_run_text(paragraph, new_text):
    """Bir paragrafin tum runlarini temizleyip ilkine yeni metni koyar (stili korur)."""
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    first_run = paragraph.runs[0]
    first_run.text = new_text
    for run in paragraph.runs[1:]:
        run.text = ''


def find_paragraph_by_prefix(prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    return None


def find_paragraph_exact(text):
    for p in doc.paragraphs:
        if p.text.strip() == text.strip():
            return p
    return None


# ============ 1. Surum satirini guncelle ============
version_p = find_paragraph_exact('Surum 1.0  ·  Mayis 2026')
if version_p is None:
    # Encoding nedeniyle muhtemelen Turkce karakterli versiyon
    for p in doc.paragraphs:
        if 'Mayıs 2026' in p.text and 'Sürüm' in p.text:
            version_p = p
            break

if version_p:
    replace_run_text(version_p, 'Sürüm 1.1  ·  Mayıs 2026 (revize)')
    print(f"OK: Surum satiri guncellendi -> {version_p.text}")
else:
    print("UYARI: Surum satiri bulunamadi")

# ============ 2. Tablo 1 (N degerleri): N=3 satirini sil ============
# Tablolar: 0=Temel Garanti, 1=N/Q tablosu, 2=Havuzlar, ...
n_table = doc.tables[1]
# Header (row 0) + 4 data row (N=3, N=5, N=7, N=11). N=3 = row 1
n_row_deleted = False
for row in list(n_table.rows):
    first_cell = row.cells[0].text.strip()
    if first_cell == '3':
        row._element.getparent().remove(row._element)
        n_row_deleted = True
        print("OK: Tablo 1'den N=3 satiri silindi")
        break
if not n_row_deleted:
    print("UYARI: Tablo 1'de N=3 satiri bulunamadi")

# ============ 3. Tablo 7 Filtre #8 guncelle ============
filter_table = doc.tables[7]
filter_updated = False
for row in filter_table.rows:
    if row.cells[0].text.strip() == '8':
        # Bu cell'in tum paragraflarini temizleyip yeni metin yaz
        cell = row.cells[2]
        # Cell'in birden fazla paragrafi olabilir; ilkini guncelle, digerlerini temizle
        if cell.paragraphs:
            replace_run_text(
                cell.paragraphs[0],
                "Admin'in son 30 günde doğrudan davet ettiği üye sayısı maksimum 1 olabilir. "
                "Topluluk yaşı < 30 gün ise limit 0'a düşer (yeni topluluk Sybil koruması — v1.1)."
            )
            for extra in cell.paragraphs[1:]:
                replace_run_text(extra, '')
        filter_updated = True
        print("OK: Tablo 7 Filtre #8 yeni topluluk kurali eklendi")
        break
if not filter_updated:
    print("UYARI: Tablo 7 Filtre #8 bulunamadi")

# ============ 4. Bolum 5.3 "Toplam ceremony suresi asla asilmaz" cumlesi ============
old_rule_p = find_paragraph_exact("Toplam ceremony süresi asla aşılmaz — yedekler ana deadline'ı kıramaz.")
if old_rule_p:
    replace_run_text(
        old_rule_p,
        "Yedek devreye girdiğinde ceremony deadline T_backup kadar uzar — state-reset stratejisi (v1.1, detay Bölüm 11)."
    )
    print("OK: Bolum 5.3 'Toplam ceremony suresi asla asilmaz' cumlesi guncellendi")
else:
    print("UYARI: Bolum 5.3 'Toplam ceremony suresi asla asilmaz' cumlesi bulunamadi")

# ============ 5. Dokuman sonuna "11. v1.1 Revizyonlari" bolumu ============
# Once "Sonuç" baslıgını bul ve oncesine ekle
sonuc_p = find_paragraph_exact("Sonuç")
if sonuc_p is None:
    # Belki farkli encoding'le yazilmis
    for p in doc.paragraphs:
        if p.text.strip() == 'Sonuç':
            sonuc_p = p
            break


def add_heading_before(anchor, text, level):
    new_p = anchor.insert_paragraph_before(text)
    # Heading stilini uygula (varsa)
    try:
        new_p.style = doc.styles[f'Heading {level}']
    except KeyError:
        pass
    return new_p


def add_paragraph_before(anchor, text, bold=False):
    new_p = anchor.insert_paragraph_before('')
    run = new_p.add_run(text)
    if bold:
        run.bold = True
    return new_p


if sonuc_p:
    add_heading_before(sonuc_p, '11. v1.1 Revizyonları', 1)

    add_paragraph_before(
        sonuc_p,
        'Bu bölüm, 2026-05-14 tarihinde dokümana yapılan revizyonları açıklar. '
        'Üç değişiklik bağımsız teknik review sonucunda tespit edilen problemleri çözer. '
        'Eski sürüme bağımlı bir referans yoksa bu bölüm gelecekte ana metne entegre edilebilir.'
    )

    # 11.1 N=3 kaldirildi
    add_heading_before(sonuc_p, '11.1 N=3 Konfigürasyonu Kaldırıldı', 2)
    add_paragraph_before(
        sonuc_p,
        'Tablo 1\'den N=3 (Q=2) satırı çıkarıldı. Desteklenen değerler artık N=5 (varsayılan), N=7 ve N=11\'dir.'
    )
    add_paragraph_before(sonuc_p, 'Gerekçe:', bold=True)
    add_paragraph_before(
        sonuc_p,
        'N=3 konfigürasyonunda Q=2 quorum ile birlikte zorunlu 1 bağımsız gözlemci kuralı koalisyon korumasını sağlamıyor: '
        'topluluk içindeki 2 üye işbirliği yaptığında Q=2 zaten tamamlanır, gözlemcinin katılımına ihtiyaç kalmaz. '
        'Yani N=3\'te "zorunlu bağımsız gözlemci" yapısı sadece tally yedekliliği sağlar, gizlilik korumasına matematiksel katkı yapmaz.'
    )
    add_paragraph_before(sonuc_p, 'Etki:', bold=True)
    add_paragraph_before(
        sonuc_p,
        'Çok küçük toplulukların (eligible üye sayısı < 5) seçim yapması zorlaşır. Bu durumda Bölüm 8.1\'deki '
        'alternatifler (eligible üyelerin rotasyonla sıralı atanması, bağımsız gözlemci sayısının artırılması) '
        'devreye girer; veya topluluk havuzu büyüyene kadar seçim ertelenir.'
    )

    # 11.2 Backup state reset
    add_heading_before(sonuc_p, '11.2 Backup Devreye Girince — State Reset ve Deadline Extend', 2)
    add_paragraph_before(
        sonuc_p,
        'Bölüm 5.3\'teki "Toplam ceremony süresi asla aşılmaz" kuralı kaldırıldı. Yeni davranış:'
    )
    add_paragraph_before(
        sonuc_p,
        '1. Bir guardian reddederse veya timeout\'a düşerse backup listesinin başındaki kişi davet edilir.'
    )
    add_paragraph_before(
        sonuc_p,
        '2. Tüm aktif guardian\'ların KEY_UPLOADED state\'i (KMP KeyCeremonyTrustee + SecureStore TrusteeJson) '
        'silinir ve PENDING\'e döner. Mobile UI "Ceremony yeniden başladı, lütfen anahtarı yeniden üret" mesajı gösterir.'
    )
    add_paragraph_before(
        sonuc_p,
        '3. Ceremony deadline yalnızca yeni gelen backup için T_backup = clamp(T_kalan / açık_slot_sayısı, 15dk, 60dk) '
        'formülü kadar uzar. Eskiden anahtar yükleyen guardian\'lar bu ek süreden faydalanmaz; '
        'orijinal T_kalan içinde geri dönmek zorundadırlar.'
    )
    add_paragraph_before(
        sonuc_p,
        '4. Substitution cap K=2: en fazla 2 backup substitution\'a izin verilir. 3. substitution gerektiğinde '
        'election otomatik CANCELLED\'a alınır ve Bulletin Board\'a MAX_SUBSTITUTIONS_REACHED kaydı yazılır.'
    )
    add_paragraph_before(sonuc_p, 'Gerekçe:', bold=True)
    add_paragraph_before(
        sonuc_p,
        'ElectionGuard key ceremony N-of-N gerektirir. KMP (votingworks/electionguard-kotlin-multiplatform) '
        'midway peer substitution\'ı desteklemez — bir guardian değiştiğinde tüm trustee\'lerin peer public keys listesi '
        'geçersiz olur ve encrypted key share exchange yeniden yapılmalıdır. Partial state korunamaz; bu yüzden '
        'state-reset zorunluluktur. Eski hazır guardian\'a ek süre vermemenin gerekçesi: guardian olmak gönüllü '
        'taahhüttür ve sistem onu tekrar bekletmek zorunda değildir. Cascade riskini K=2 cap sınırlar.'
    )
    add_paragraph_before(sonuc_p, 'Etkilenen alanlar:', bold=True)
    add_paragraph_before(
        sonuc_p,
        '• Bölüm 5.3 (Adaptif Yedek Atama) — yeni davranış burada okunur.\n'
        '• Tablo 14 (örnek senaryolar) — geçerli kalır, hesaplama formülü değişmedi.\n'
        '• elections tablosu — yeni kolon substitution_count INT DEFAULT 0 eklenir.\n'
        '• election_guardians.status enum\'una RESET geçiş state\'i eklenir (PENDING\'e dönüş sırasında).'
    )

    # 11.3 Yeni topluluk Sybil korumasi
    add_heading_before(sonuc_p, '11.3 Yeni Topluluk Sybil Koruması', 2)
    add_paragraph_before(
        sonuc_p,
        'Tablo 7 Filtre #8 (Admin\'in Doğrudan Davetlisi Limiti) genişletildi. '
        'Topluluk yaşı < 30 gün ise admin\'in son 30 günde davet ettiği maksimum 1 üye sınırı 0\'a düşer. '
        'Yani yeni topluluğun ilk 30 gününde admin\'in davet ettiği hiç kimse guardian olamaz.'
    )
    add_paragraph_before(sonuc_p, 'Gerekçe:', bold=True)
    add_paragraph_before(
        sonuc_p,
        'Bölüm 2.1\'deki "topluluk yaşı istisnası" (topluluk yaşı < 14 gün ise üyelik yaşı şartı uygulanmaz) ile '
        'admin manuel mod kombinasyonunda Sybil deliği oluşur: admin yeni topluluk açar, 4 yandaşını davet eder, '
        'hepsini guardian yapar. Manuel modu tamamen yasaklamak topluluk içi sosyal seçimi de engeller; '
        'doğrudan davet limitini 0\'a indirmek manuel modu çalışır tutarken admin\'in nüfuz alanını kapatır.'
    )
    add_paragraph_before(sonuc_p, 'Etki:', bold=True)
    add_paragraph_before(
        sonuc_p,
        'Yeni topluluğun ilk 30 gününde guardian havuzu genellikle dar olacaktır (sadece kendi başına opt-in yapmış '
        'üyeler ve admin\'in davet etmediği başka kanaldan gelmiş üyeler). Havuz < N olursa Bölüm 8.1\'deki '
        'küçük topluluk seçenekleri devreye girer; veya seçim 30 günlük topluluk yaşı eşiği geçene kadar ertelenir.'
    )

    # 11.4 Karar Ozeti tablosu
    add_heading_before(sonuc_p, '11.4 Karar Özeti', 2)
    add_paragraph_before(
        sonuc_p,
        'Üç değişiklik dokümanın temel felsefesini değiştirmez: hibrit manuel atama + opsiyonel sortition + '
        'zorunlu bağımsız gözlemci + post-hoc denetim katmanları korunur. Bu revizyonlar yapısal koruma '
        'eksiklerini matematiksel olarak kapatır ve KMP implementation gerçekliğiyle uyumlu hale getirir.'
    )

    print("OK: v1.1 Revizyonlari bolumu eklendi (4 alt baslik)")
else:
    print("UYARI: Sonuc bolumu bulunamadi, revizyon bolumu eklenemedi")

# ============ Kaydet ============
doc.save(DOC_PATH)
print(f"\nDokuman kaydedildi: {DOC_PATH}")
print("Yedek: CepSandik_Guardian_Secim_Sureci.v1.0.backup.docx")
