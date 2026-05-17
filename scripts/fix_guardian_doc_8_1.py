"""
v1.1 sonrasi kalan tek celiski duzeltmesi:
Bolum 8.1 (Cok Kucuk Topluluk) hala "N=3, Q=2" oneriyor — N=3 v1.1'de kaldirildi.
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

DOC = 'CepSandik_Guardian_Secim_Sureci.docx'
doc = Document(DOC)


def replace_run_text(p, txt):
    if not p.runs:
        p.add_run(txt); return
    p.runs[0].text = txt
    for r in p.runs[1:]:
        r.text = ''


fixed = False
for p in doc.paragraphs:
    if p.text.strip() == 'N değerini düşürme önerisi (N=3, Q=2)':
        replace_run_text(
            p,
            'Minimum N=5 (Q=3) ile devam (eligible havuz ≥ 5 ise); '
            'havuz yetersizse seçim, topluluk havuzu büyüyene kadar ertelenir (v1.1: N=3 kaldırıldı)'
        )
        fixed = True
        print('OK: Bolum 8.1 N=3 onerisi N=5 minimum ile degistirildi')
        break

if not fixed:
    print('UYARI: Bolum 8.1 N=3 cumlesi bulunamadi')

doc.save(DOC)
print(f'Kaydedildi: {DOC}')
